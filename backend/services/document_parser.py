"""
Document parsing service for Chronicle
Supports PDF, DOCX, MD, and TXT files with full-text extraction
"""

import os
import tempfile
from pathlib import Path
from typing import Dict, Optional
import logging

try:
    import PyPDF2
except ImportError:
    import pypdf2 as PyPDF2

try:
    from docx import Document
except ImportError:
    Document = None

import markdown

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parse various document formats for searchable text content"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.md', '.txt'}
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file type is supported"""
        return Path(filename).suffix.lower() in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def parse_document(cls, file_path: str, filename: str) -> Dict[str, str]:
        """
        Parse document and extract text content
        
        Returns:
            Dict with keys: content, file_type, page_count, word_count
        """
        try:
            extension = Path(filename).suffix.lower()
            
            if extension == '.pdf':
                return cls._parse_pdf(file_path, filename)
            elif extension == '.docx':
                return cls._parse_docx(file_path, filename)
            elif extension == '.md':
                return cls._parse_markdown(file_path, filename)
            elif extension == '.txt':
                return cls._parse_text(file_path, filename)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
                
        except Exception as e:
            logger.error(f"Error parsing {filename}: {str(e)}")
            return {
                'content': f"Error parsing document: {str(e)}",
                'file_type': extension,
                'page_count': 0,
                'word_count': 0,
                'error': str(e)
            }
    
    @classmethod
    def _parse_pdf(cls, file_path: str, filename: str) -> Dict[str, str]:
        """Parse PDF document"""
        content = []
        page_count = 0
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1} of {filename}: {str(e)}")
                    content.append(f"[Page {page_num + 1}]\n[Text extraction failed]")
        
        full_content = '\n\n'.join(content)
        word_count = len(full_content.split())
        
        return {
            'content': full_content,
            'file_type': 'pdf',
            'page_count': page_count,
            'word_count': word_count
        }
    
    @classmethod
    def _parse_docx(cls, file_path: str, filename: str) -> Dict[str, str]:
        """Parse DOCX document"""
        if Document is None:
            raise ImportError("python-docx not available")
        
        doc = Document(file_path)
        content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    content.append(row_text)
        
        full_content = '\n'.join(content)
        word_count = len(full_content.split())
        
        return {
            'content': full_content,
            'file_type': 'docx',
            'page_count': 1,  # DOCX doesn't have clear page breaks
            'word_count': word_count
        }
    
    @classmethod
    def _parse_markdown(cls, file_path: str, filename: str) -> Dict[str, str]:
        """Parse Markdown document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Convert markdown to plain text (remove markdown syntax)
        html = markdown.markdown(content)
        # Simple HTML tag removal for plain text
        import re
        plain_text = re.sub(r'<[^>]+>', '', html)
        
        word_count = len(plain_text.split())
        
        return {
            'content': plain_text,
            'file_type': 'markdown',
            'page_count': 1,
            'word_count': word_count,
            'raw_markdown': content  # Keep original for display
        }
    
    @classmethod
    def _parse_text(cls, file_path: str, filename: str) -> Dict[str, str]:
        """Parse plain text document"""
        encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'latin1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError(f"Could not decode text file {filename} with any supported encoding")
        
        word_count = len(content.split())
        
        return {
            'content': content,
            'file_type': 'text',
            'page_count': 1,
            'word_count': word_count
        }

# Convenience function for the API
def parse_uploaded_document(file_content: bytes, filename: str) -> Dict[str, str]:
    """
    Parse an uploaded document from bytes
    
    Args:
        file_content: Raw file bytes
        filename: Original filename (for type detection)
    
    Returns:
        Parsed document data
    """
    if not DocumentParser.is_supported(filename):
        raise ValueError(f"Unsupported file type: {Path(filename).suffix}")
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name
    
    try:
        return DocumentParser.parse_document(temp_path, filename)
    finally:
        # Clean up temporary file
        os.unlink(temp_path)